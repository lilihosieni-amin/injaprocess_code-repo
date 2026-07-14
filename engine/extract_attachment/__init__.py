from engine_common import data_root


def attachments_dir(root, dept):
    return root / "departments" / dept / "attachments"


def text_dir(root, dept):
    return attachments_dir(root, dept) / ".text"


def find_docx(root, dept):
    adir = attachments_dir(root, dept)
    # glob on a missing directory yields nothing; .text/ is a subdir so *.docx
    # at this level never descends into it.
    return sorted(p for p in adir.glob("*.docx") if p.is_file())


def docx_to_text(path):
    from docx import Document  # lazy — keeps import cost out of the fast paths
    doc = Document(str(path))
    lines = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                lines.append(cell.text)
    return "\n".join(lines).strip() + "\n"


def needs_conversion(src, dst):
    return (not dst.exists()) or (src.stat().st_mtime > dst.stat().st_mtime)


def run_extract_attachment(dept, root=None, convert=None):
    root = root or data_root()
    convert = convert or docx_to_text
    tdir = text_dir(root, dept)
    ok, errors = [], []
    for src in find_docx(root, dept):
        dst = tdir / (src.stem + ".txt")
        try:
            if needs_conversion(src, dst):
                tdir.mkdir(parents=True, exist_ok=True)
                dst.write_text(convert(src), encoding="utf-8")
            ok.append(dst.relative_to(root).as_posix())
        except Exception as e:  # one bad doc must not sink the rest (supplement)
            errors.append((src.name, str(e)))
    return ok, errors
